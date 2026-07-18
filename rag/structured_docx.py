"""Word 文档结构化切分（按章节 / Q&A / 关键词列表）。"""

from __future__ import annotations

import hashlib
import re
from pathlib import Path

from docx import Document
from langchain_core.documents import Document as LCDocument

SECTION_FAQ = "一、常见问题（回复）"
SECTION_UPGRADE = "触发客户合作关键信息"
SECTION_REJECT = "不合作的原因"

CHUNK_FAQ = "faq_qa"
CHUNK_UPGRADE = "upgrade_keyword"
CHUNK_REJECT = "reject_reason"
CHUNK_SECTION = "section_header"

_NUMBERED_Q = re.compile(r"^\d+[、.]")
_ANSWER_WRAP = re.compile(r"^（.*）$")


def load_structured_docx(file_path: str | Path) -> list[LCDocument]:
    """加载 docx 并按业务结构切分为 chunk。"""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文档不存在: {path}")

    paragraphs = [p.text.strip() for p in Document(str(path)).paragraphs if p.text.strip()]
    source = path.name
    chunks = _split_bd_prompt_paragraphs(paragraphs, source=source)

    if not chunks:
        raise ValueError(f"未能从文档切分出有效 chunk: {path}")

    return chunks


def _split_bd_prompt_paragraphs(paragraphs: list[str], source: str) -> list[LCDocument]:
    chunks: list[LCDocument] = []
    chunk_index = 0
    current_section = ""
    i = 0

    while i < len(paragraphs):
        text = paragraphs[i]

        if text in (SECTION_FAQ, SECTION_UPGRADE, SECTION_REJECT):
            current_section = text
            chunks.append(
                _make_chunk(
                    content=f"章节：{text}",
                    source=source,
                    chunk_index=chunk_index,
                    section=current_section,
                    chunk_type=CHUNK_SECTION,
                )
            )
            chunk_index += 1
            i += 1
            continue

        if current_section == SECTION_FAQ:
            consumed, faq_chunks = _parse_faq_block(paragraphs, i, source, chunk_index, current_section)
            chunks.extend(faq_chunks)
            chunk_index += len(faq_chunks)
            i += consumed
            continue

        if current_section == SECTION_UPGRADE:
            chunks.append(
                _make_chunk(
                    content=f"触发合作关键词：{text}",
                    source=source,
                    chunk_index=chunk_index,
                    section=current_section,
                    chunk_type=CHUNK_UPGRADE,
                    keyword=text,
                )
            )
            chunk_index += 1
            i += 1
            continue

        if current_section == SECTION_REJECT:
            chunks.append(
                _make_chunk(
                    content=f"不合作原因：{text}",
                    source=source,
                    chunk_index=chunk_index,
                    section=current_section,
                    chunk_type=CHUNK_REJECT,
                    reason=text,
                )
            )
            chunk_index += 1
            i += 1
            continue

        i += 1

    return chunks


def _parse_faq_block(
    paragraphs: list[str],
    start: int,
    source: str,
    start_index: int,
    section: str,
) -> tuple[int, list[LCDocument]]:
    """解析 FAQ：问题 + （回答）或下一条非括号回答。"""
    chunks: list[LCDocument] = []
    i = start
    idx = start_index

    while i < len(paragraphs):
        text = paragraphs[i]
        if text in (SECTION_UPGRADE, SECTION_REJECT):
            break

        if _ANSWER_WRAP.match(text):
            i += 1
            continue

        question = text
        answer = ""
        i += 1

        if i < len(paragraphs) and _ANSWER_WRAP.match(paragraphs[i]):
            answer = paragraphs[i].strip("（）")
            i += 1
        elif i < len(paragraphs) and paragraphs[i] not in (SECTION_UPGRADE, SECTION_REJECT):
            if not _NUMBERED_Q.match(paragraphs[i]) and not _ANSWER_WRAP.match(paragraphs[i]):
                answer = paragraphs[i]
                i += 1

        content = f"问题：{question}\n回答：{answer or '（见知识库原文）'}"
        chunks.append(
            _make_chunk(
                content=content,
                source=source,
                chunk_index=idx,
                section=section,
                chunk_type=CHUNK_FAQ,
                question=question,
            )
        )
        idx += 1

    return i - start, chunks


def _make_chunk(
    content: str,
    source: str,
    chunk_index: int,
    section: str,
    chunk_type: str,
    question: str = "",
    keyword: str = "",
    reason: str = "",
) -> LCDocument:
    chunk_id = _stable_chunk_id(source, section, chunk_type, chunk_index, content[:40])
    metadata = {
        "source": source,
        "chunk_id": chunk_id,
        "chunk_index": chunk_index,
        "page": 0,
        "section": section,
        "chunk_type": chunk_type,
    }
    if question:
        metadata["question"] = question[:200]
    if keyword:
        metadata["keyword"] = keyword[:200]
    if reason:
        metadata["reason"] = reason[:200]

    return LCDocument(page_content=content, metadata=metadata)


def _stable_chunk_id(
    source: str,
    section: str,
    chunk_type: str,
    chunk_index: int,
    prefix: str,
) -> str:
    raw = f"{source}|{section}|{chunk_type}|{chunk_index}|{prefix}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]
